"""
Module de Scraping Tanmia.ma - MVP v1.8
Extraction robuste avec PARSING COMPLET des fichiers attachés (PDF, DOC, DOCX)

CHANGELOG v1.8:
- NOUVEAU: download_and_parse_attachment() - Télécharge et extrait le texte
- NOUVEAU: Support PDF (pdfplumber), DOCX (python-docx), DOC (antiword fallback)
- NOUVEAU: Champ 'contenu_texte' dans fichiers_attaches
- NOUVEAU: Extraction emails depuis contenu fichiers
- Limite: 5000 chars/fichier pour éviter surcharge IA
"""
import requests
from bs4 import BeautifulSoup
import time
import random
import re
import os
import tempfile
from typing import List, Dict, Optional, Set, Tuple
from urllib.parse import urljoin
from io import BytesIO


# ============================================================================
# CONFIGURATION
# ============================================================================

BASE_URL = "https://tanmia.ma"
USER_AGENTS = [
    'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
    'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
]

TIMEOUT = 30
DOWNLOAD_TIMEOUT = 60  # Plus long pour fichiers volumineux

# Extensions de fichiers à détecter
VALID_FILE_EXTENSIONS = [
    '.pdf', '.doc', '.docx', 
    '.xls', '.xlsx', 
    '.ppt', '.pptx',
    '.zip', '.rar',
    '.odt', '.ods', '.odp'
]

# Extensions parsables (v1.8)
PARSABLE_EXTENSIONS = ['.pdf', '.doc', '.docx']

# Limite contenu fichier pour IA
MAX_FILE_CONTENT_LENGTH = 5000


# ============================================================================
# UTILITAIRES
# ============================================================================

def get_random_headers() -> Dict[str, str]:
    """Génère des headers HTTP aléatoires."""
    return {
        'User-Agent': random.choice(USER_AGENTS),
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
        'Accept-Language': 'fr-FR,fr;q=0.9',
    }


def human_delay(min_sec: float = 2.0, max_sec: float = 5.0) -> None:
    """Délai aléatoire pour simuler comportement humain."""
    time.sleep(random.uniform(min_sec, max_sec))


def create_session() -> requests.Session:
    """Crée une session HTTP avec headers configurés."""
    session = requests.Session()
    session.headers.update(get_random_headers())
    return session


# ============================================================================
# EXTRACTION EMAILS DEPUIS TEXTE (v1.8)
# ============================================================================

def extract_emails_from_text(text: str) -> List[str]:
    """
    Extrait tous les emails d'un texte.
    
    Supporte formats:
    - Standard: email@domain.com
    - Espaces: email @ domain . com
    - Obfusqué: email AT domain DOT com
    
    Args:
        text: Texte à analyser
    
    Returns:
        Liste d'emails uniques
    """
    if not text:
        return []
    
    emails = set()
    
    # Pattern standard
    pattern1 = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    emails.update(re.findall(pattern1, text))
    
    # Pattern avec espaces
    pattern2 = r'\b([A-Za-z0-9._%+-]+)\s*@\s*([A-Za-z0-9.-]+)\s*\.\s*([A-Z|a-z]{2,})\b'
    for match in re.finditer(pattern2, text):
        email = f"{match.group(1)}@{match.group(2)}.{match.group(3)}"
        emails.add(email)
    
    # Pattern AT/DOT
    pattern3 = r'\b([A-Za-z0-9._%+-]+)\s+(?:at|AT)\s+([A-Za-z0-9.-]+)\s+(?:dot|DOT)\s+([A-Z|a-z]{2,})\b'
    for match in re.finditer(pattern3, text):
        email = f"{match.group(1)}@{match.group(2)}.{match.group(3)}"
        emails.add(email)
    
    # Pattern [at] [dot]
    pattern4 = r'\b([A-Za-z0-9._%+-]+)\s*\[at\]\s*([A-Za-z0-9.-]+)\s*\[dot\]\s*([A-Z|a-z]{2,})\b'
    for match in re.finditer(pattern4, text, re.IGNORECASE):
        email = f"{match.group(1)}@{match.group(2)}.{match.group(3)}"
        emails.add(email)
    
    # Nettoyage
    cleaned = []
    for email in emails:
        email = re.sub(r'\s+', '', email).lower()
        if '@' in email and '.' in email.split('@')[1]:
            cleaned.append(email)
    
    return list(set(cleaned))


# ============================================================================
# PARSING FICHIERS (NOUVEAU v1.8)
# ============================================================================

def parse_pdf_content(file_bytes: bytes) -> Tuple[str, List[str]]:
    """Parse un fichier PDF."""
    try:
        import pdfplumber
        
        text_parts = []
        
        with pdfplumber.open(BytesIO(file_bytes)) as pdf:
            for page in pdf.pages[:20]:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        
        full_text = '\n'.join(text_parts)
        
        # CORRIGÉ: Extraire emails sur TOUT le texte AVANT troncature
        emails = extract_emails_from_text(full_text)
        
        # Tronquer seulement pour l'IA après extraction
        if len(full_text) > MAX_FILE_CONTENT_LENGTH:
            full_text = full_text[:MAX_FILE_CONTENT_LENGTH] + "...[tronqué]"
        
        return full_text, emails
        
    except ImportError:
        print("⚠️ pdfplumber non installé")
        return "", []
    except Exception as e:
        print(f"⚠️ Erreur parsing PDF: {e}")
        return "", []


def parse_docx_content(file_bytes: bytes) -> Tuple[str, List[str]]:
    """Parse un fichier DOCX."""
    try:
        from docx import Document
        
        doc = Document(BytesIO(file_bytes))
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        full_text = '\n'.join(text_parts)
        
        # CORRIGÉ: Extraire emails AVANT troncature
        emails = extract_emails_from_text(full_text)
        
        if len(full_text) > MAX_FILE_CONTENT_LENGTH:
            full_text = full_text[:MAX_FILE_CONTENT_LENGTH] + "...[tronqué]"
        
        return full_text, emails
        
    except ImportError:
        print("⚠️ python-docx non installé")
        return "", []
    except Exception as e:
        print(f"⚠️ Erreur parsing DOCX: {e}")
        return "", []


def parse_doc_content(file_bytes: bytes) -> Tuple[str, List[str]]:
    """Parse un fichier DOC (ancien format)."""
    try:
        import subprocess
        import tempfile
        import os
        
        with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name
        
        full_text = ""
        
        # Tenter antiword
        try:
            result = subprocess.run(['antiword', tmp_path], capture_output=True, text=True, timeout=30)
            if result.returncode == 0:
                full_text = result.stdout
        except FileNotFoundError:
            print("⚠️ antiword/catdoc non installé (apt install antiword)")
        
        os.unlink(tmp_path)
        
        if full_text:
            emails = extract_emails_from_text(full_text)
            if len(full_text) > MAX_FILE_CONTENT_LENGTH:
                full_text = full_text[:MAX_FILE_CONTENT_LENGTH] + "...[tronqué]"
            return full_text, emails
        
        return "", []
        
    except Exception as e:
        print(f"⚠️ Erreur parsing DOC: {e}")
        return "", []

def download_and_parse_attachment(
    url: str, 
    file_type: str, 
    session: Optional[requests.Session] = None
) -> Tuple[str, List[str]]:
    """
    Télécharge et parse un fichier attaché.
    
    Args:
        url: URL du fichier
        file_type: Type de fichier (pdf, doc, docx)
        session: Session HTTP
    
    Returns:
        Tuple (contenu_texte, emails_extraits)
    """
    if session is None:
        session = create_session()
    
    # Vérifier si parsable
    if file_type not in ['pdf', 'doc', 'docx']:
        return "", []
    
    try:
        print(f"      📥 Téléchargement {file_type.upper()}...")
        
        # Télécharger
        response = session.get(url, timeout=DOWNLOAD_TIMEOUT, stream=True)
        response.raise_for_status()
        
        # Vérifier taille (max 10MB)
        content_length = response.headers.get('Content-Length')
        if content_length and int(content_length) > 10 * 1024 * 1024:
            print(f"      ⚠️ Fichier trop volumineux (>10MB), ignoré")
            return "", []
        
        file_bytes = response.content
        
        # Parser selon type
        if file_type == 'pdf':
            return parse_pdf_content(file_bytes)
        elif file_type == 'docx':
            return parse_docx_content(file_bytes)
        elif file_type == 'doc':
            return parse_doc_content(file_bytes)
        else:
            return "", []
            
    except requests.RequestException as e:
        print(f"      ⚠️ Erreur téléchargement: {e}")
        return "", []
    except Exception as e:
        print(f"      ⚠️ Erreur parsing: {e}")
        return "", []


# ============================================================================
# EXTRACTION FICHIERS ATTACHÉS (v1.8 enrichi)
# ============================================================================

def extract_attachments(
    soup: BeautifulSoup, 
    session: Optional[requests.Session] = None,
    parse_content: bool = True
) -> List[Dict]:
    """
    Extrait les fichiers attachés avec parsing du contenu (v1.8).
    
    Args:
        soup: Objet BeautifulSoup de la page
        session: Session HTTP pour téléchargement
        parse_content: Si True, télécharge et parse les fichiers PDF/DOC/DOCX
    
    Returns:
        Liste de dicts avec clés: nom, url, type, contenu_texte, emails_fichier
    """
    fichiers: List[Dict] = []
    seen_urls: Set[str] = set()
    
    # Sélecteurs multiples
    selectors = [
        'ul.post-attachments a',
        'ul.wp-block-file a',
        'div.wp-block-file a',
        'a.attachment-link',
        'a.wp-block-file__button',
        'div.elementor-widget-attachment a',
        'div.elementor-widget-icon-list a',
        'a.download-link',
        'a[download]',
    ]
    
    # Ajouter sélecteurs par extension
    for ext in VALID_FILE_EXTENSIONS:
        selectors.append(f'a[href$="{ext}"]')
        selectors.append(f'a[href$="{ext.upper()}"]')
    
    # Extraction
    for selector in selectors:
        try:
            for a_tag in soup.select(selector):
                href = a_tag.get('href', '')
                
                if not href or href in seen_urls:
                    continue
                
                if href.startswith('#') or href.startswith('javascript:'):
                    continue
                
                is_file = any(ext in href.lower() for ext in VALID_FILE_EXTENSIONS)
                if not is_file:
                    continue
                
                seen_urls.add(href)
                
                # Nom du fichier
                nom = a_tag.get_text(strip=True)
                if not nom or len(nom) < 3:
                    nom = href.split('/')[-1].split('?')[0]
                    try:
                        from urllib.parse import unquote
                        nom = unquote(nom)
                    except:
                        pass
                
                nom = nom[:100].strip()
                
                # Type de fichier
                type_fichier = 'autre'
                href_lower = href.lower()
                
                for ext in VALID_FILE_EXTENSIONS:
                    if ext in href_lower:
                        type_fichier = ext.replace('.', '')
                        break
                
                # URL absolue
                url_absolue = urljoin(BASE_URL, href)
                
                # Structure de base
                fichier_data = {
                    'nom': nom,
                    'url': url_absolue,
                    'type': type_fichier,
                    'contenu_texte': '',      # NOUVEAU v1.8
                    'emails_fichier': []       # NOUVEAU v1.8
                }
                
                # ============================================================
                # PARSING CONTENU (NOUVEAU v1.8)
                # ============================================================
                if parse_content and type_fichier in ['pdf', 'doc', 'docx']:
                    contenu, emails = download_and_parse_attachment(
                        url_absolue, 
                        type_fichier, 
                        session
                    )
                    fichier_data['contenu_texte'] = contenu
                    fichier_data['emails_fichier'] = emails
                    
                    if contenu:
                        print(f"      ✅ Parsé: {len(contenu)} chars, {len(emails)} email(s)")
                
                fichiers.append(fichier_data)
                
        except Exception as e:
            print(f"⚠️ Erreur sélecteur {selector}: {e}")
            continue
    
    return fichiers


# ============================================================================
# SCRAPING - PAGE LISTING
# ============================================================================

def scrape_listing_page(url: str, session: Optional[requests.Session] = None) -> List[str]:
    """Scrape une page de listing pour extraire les URLs."""
    if session is None:
        session = create_session()
    
    try:
        human_delay(2, 4)
        response = session.get(url, timeout=TIMEOUT)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, 'lxml')
        
        urls: Set[str] = set()
        articles = soup.select('article.elementor-post')
        
        for article in articles:
            title_link = article.select_one('h3.elementor-post__title a')
            if title_link and title_link.get('href'):
                href = title_link['href']
                if '/evenement/' not in href:
                    urls.add(href)
        
        return list(urls)
    
    except Exception as e:
        print(f"❌ Erreur scraping listing: {e}")
        return []


# ============================================================================
# SCRAPING - PAGE DÉTAIL (v1.8 avec parsing fichiers)
# ============================================================================

def scrape_detail_page(
    url: str, 
    session: Optional[requests.Session] = None,
    parse_attachments: bool = True
) -> Optional[Dict]:
    """
    Scrape une page de détail avec parsing complet des fichiers attachés.
    
    Args:
        url: URL de la page
        session: Session HTTP
        parse_attachments: Si True, télécharge et parse les PDF/DOC/DOCX
    
    Returns:
        Dict enrichi avec contenu des fichiers et emails extraits
    """
    if session is None:
        session = create_session()
    
    try:
        human_delay(1, 3)
        response = session.get(url, timeout=TIMEOUT)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, 'lxml')
        
        # Titre
        titre = "Non spécifié"
        for sel in ['h1.elementor-heading-title', 'h1.entry-title', 'h1']:
            elem = soup.select_one(sel)
            if elem and elem.get_text(strip=True):
                titre = elem.get_text(strip=True)
                break
        
        # Organisation
        organisation = "À déterminer par IA"
        
        # Date
        date = "Non spécifié"
        for sel in ['time', 'span.elementor-post-info__item--type-date', '.elementor-post-date', '.entry-date']:
            date_elem = soup.select_one(sel)
            if date_elem:
                text = date_elem.get_text(strip=True)
                if any(char.isdigit() for char in text) and len(text) < 50:
                    date = text
                    break
        
        # ====================================================================
        # FICHIERS ATTACHÉS AVEC PARSING (v1.8)
        # ====================================================================
        print(f"   📎 Extraction fichiers attachés...")
        fichiers_attaches = extract_attachments(soup, session, parse_attachments)
        
        # Collecter tous les emails des fichiers
        emails_from_files = []
        for f in fichiers_attaches:
            emails_from_files.extend(f.get('emails_fichier', []))
        emails_from_files = list(set(emails_from_files))
        
        # ====================================================================
        # ZONE DE CONTENU
        # ====================================================================
        content_zone = None
        for sel in ['div.elementor-widget-theme-post-content', 'article.elementor-post', 'div.entry-content', 'main']:
            zone = soup.select_one(sel)
            if zone:
                content_zone = zone
                break
        
        # Nettoyage
        if content_zone:
            for tag in content_zone.select('nav, header, footer, aside, script, style, iframe, .breadcrumbs, .share-buttons, .post-navigation, ul.post-attachments, .elementor-widget-shortcode, .elementor-social-icons-wrapper'):
                tag.decompose()
        else:
            content_zone = soup
            for tag in soup.select('nav, header, footer, aside, script, style, iframe'):
                tag.decompose()
        
        # Texte complet
        texte_complet = content_zone.get_text(separator='\n', strip=True)
        texte_complet = re.sub(r'\n\s*\n', '\n\n', texte_complet)
        texte_complet = re.sub(r' +', ' ', texte_complet)
        texte_complet = texte_complet.strip()
        
        if len(texte_complet) > 12000:
            texte_complet = texte_complet[:12000] + "...[texte tronqué]"
        
        # ====================================================================
        # RETOUR ENRICHI (v1.8)
        # ====================================================================
        return {
            'url': url,
            'titre': titre,
            'organisation': organisation,
            'date': date,
            'texte_complet': texte_complet,
            'fichiers_attaches': fichiers_attaches,      # Avec contenu_texte
            'emails_from_files': emails_from_files       # NOUVEAU v1.8
        }
    
    except Exception as e:
        print(f"❌ Erreur scraping détail: {e}")
        return None


# ============================================================================
# SCRAPING - FONCTION PRINCIPALE
# ============================================================================

def scrape_tanmia(
    url_type: str, 
    max_pages: int, 
    progress_callback=None,
    parse_attachments: bool = True
) -> List[Dict]:
    """
    Fonction principale de scraping avec parsing fichiers (v1.8).
    
    Args:
        url_type: "appels-doffres" ou "offres-demploi"
        max_pages: Nombre de pages à scraper
        progress_callback: Callback progression
        parse_attachments: Si True, parse le contenu des PDF/DOC/DOCX
    
    Returns:
        Liste de dicts avec données complètes
    """
    session = create_session()
    all_urls: List[str] = []
    results: List[Dict] = []
    
    # Phase 1: Listings
    print(f"\n🔍 PHASE 1: Scraping des listings ({max_pages} pages)")
    print("=" * 60)
    
    for page_num in range(1, max_pages + 1):
        page_url = f"{BASE_URL}/{url_type}/" if page_num == 1 else f"{BASE_URL}/{url_type}/{page_num}/"
        print(f"📄 Page {page_num}/{max_pages}: {page_url}")
        
        urls = scrape_listing_page(page_url, session)
        print(f"   ✅ {len(urls)} opportunités trouvées")
        all_urls.extend(urls)
        
        if progress_callback:
            progress_callback(page_num / max_pages * 0.3, f"Page {page_num}/{max_pages}")
    
    all_urls = list(set(all_urls))
    print(f"\n✅ Total unique: {len(all_urls)} opportunités")
    
    # Phase 2: Détails + Parsing
    print(f"\n📊 PHASE 2: Extraction détails + PARSING FICHIERS (v1.8)")
    print("=" * 60)
    
    total_fichiers = 0
    total_emails_fichiers = 0
    
    for idx, url in enumerate(all_urls, 1):
        print(f"\n🔗 {idx}/{len(all_urls)}: {url[:60]}...")
        
        detail = scrape_detail_page(url, session, parse_attachments)
        
        if detail:
            nb_fichiers = len(detail.get('fichiers_attaches', []))
            nb_emails_files = len(detail.get('emails_from_files', []))
            total_fichiers += nb_fichiers
            total_emails_fichiers += nb_emails_files
            
            # Compter fichiers parsés
            nb_parses = sum(1 for f in detail.get('fichiers_attaches', []) if f.get('contenu_texte'))
            
            print(f"   ✅ {nb_fichiers} fichier(s) | {nb_parses} parsé(s) | {nb_emails_files} email(s) extraits")
            results.append(detail)
        else:
            print(f"   ❌ Échec")
        
        if progress_callback:
            progress = 0.3 + (idx / len(all_urls)) * 0.7
            progress_callback(progress, f"Détail {idx}/{len(all_urls)}")
    
    # Résumé
    print(f"\n{'=' * 60}")
    print(f"✅ SCRAPING v1.8 TERMINÉ")
    print(f"   • Opportunités: {len(results)}/{len(all_urls)}")
    print(f"   • Fichiers détectés: {total_fichiers}")
    print(f"   • Emails extraits des fichiers: {total_emails_fichiers}")
    print(f"{'=' * 60}")
    
    return results


# ============================================================================
# TEST
# ============================================================================

def test_scraper():
    """Test du scraper v1.8."""
    print("🧪 TEST SCRAPER v1.8 (avec parsing fichiers)")
    print("=" * 60)
    
    # Test listing
    print("\n📄 Test 1: Scraping listing...")
    urls = scrape_listing_page("https://tanmia.ma/appels-doffres/")
    print(f"   URLs trouvées: {len(urls)}")
    
    if urls:
        print(f"\n📎 Test 2: Extraction + parsing fichiers...")
        print(f"   URL: {urls[0]}")
        
        detail = scrape_detail_page(urls[0], parse_attachments=True)
        
        if detail:
            print(f"\n   ✅ Résultat:")
            print(f"      Titre: {detail['titre'][:50]}...")
            
            fichiers = detail.get('fichiers_attaches', [])
            print(f"\n   📎 Fichiers: {len(fichiers)}")
            
            for f in fichiers:
                print(f"      • {f['nom']} ({f['type']})")
                if f.get('contenu_texte'):
                    print(f"        Contenu: {len(f['contenu_texte'])} chars")
                if f.get('emails_fichier'):
                    print(f"        Emails: {f['emails_fichier']}")
            
            print(f"\n   📧 Emails from files: {detail.get('emails_from_files', [])}")
    
    print("\n✅ Test terminé")


if __name__ == "__main__":
    test_scraper()
